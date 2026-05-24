from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("D:/MagicMod")
DOCX_PATH = ROOT / "docs" / "MODEL_REPLACEMENT_TEST_REPORT.docx"
SCREENSHOT_PATH = ROOT / "docs" / "model_replacement_stage2_screenshot.png"


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(8.5)


def add_command_case(document, step, title, command, expected, result):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"{step}. {title}")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_code_block(document, [command])
    add_body(document, f"预期效果：{expected}；实测结果：{result}")


def build():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MagicMod 模型替换功能测试报告")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Forge 1.21.11 | 客户端实体模型替换、Bedrock 动画、FBX 演示、MMESH 中间格式")
    sub_run.font.name = "Microsoft YaHei"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(90, 90, 90)

    add_heading(document, "1. 验证范围", 1)
    add_body(document, "本报告记录本轮模型替换功能的代码审计、开发补充、真实游戏内验证和资源释放检查。测试在 Minecraft Forge 1.21.11 单人世界中完成。")
    add_bullets(document, [
        "验证精确 target 实体替换，而不是只替换整个实体类型。",
        "验证 Bedrock geo.json 与 animation.json 的加载、渲染和动画播放。",
        "验证 ASCII FBX 演示模型接入统一 runtime，并可播放 fallback 动画。",
        "验证新增 .mmesh 中间格式可加载、渲染并播放 fallback 动画。",
        "验证 reload 后 runtimeInstances、loadedModels、liveGpuBuffers 回落到 0。",
    ])

    add_heading(document, "2. 代码审计结论", 1)
    add_bullets(document, [
        "修复 replaceType(Identifier, Identifier) 忽略 modelId 的问题。",
        "新增 ModelReplacementTarget，支持 entity_type、精确实体、uuid、player_name、tag。",
        "将 /magicmodel client replace target 修正为替换当前目标实体本身。",
        "新增 MMESH loader 和测试资源，作为 FBX 离线转换后的运行时中间格式。",
        "清理旧版未使用命令处理函数和未使用格式建议常量。",
        "修复实体级 runtime 实例剪枝前缀不一致的问题，避免实体消失后残留运行时实例。",
        "保留对完整服务端同步与二进制 FBX 骨骼动画的边界说明，避免把未完成能力伪装为已完成。",
    ])

    add_heading(document, "3. 指令测试矩阵", 1)
    command_cases = [
        ("测试立方体精确替换", "/magicmodel client replace target magicmod:test_cube", "精确替换当前玩家实体为测试立方体", "通过，渲染计数 0 -> 1386"),
        ("循环播放基础动画", "/magicmodel client animation play target spin loop", "循环播放 spin 动画", "通过"),
        ("播放一次性动画", "/magicmodel client animation play target pulse once", "播放一次 pulse 动画", "通过"),
        ("清理当前目标替换", "/magicmodel client clear target", "清理当前实体替换", "通过"),
        ("Bedrock 模型替换", "/magicmodel client replace target bedrock magicmod:entity_models/bedrock/test_creature.geo.json minecraft:textures/block/copper_block.png magicmod:entity_models/bedrock/test_creature.animation.json", "替换为 Bedrock 模型并绑定动画文件", "通过，渲染计数 2794 -> 4198"),
        ("Bedrock 指定动画播放", "/magicmodel client animation play target animation.magicmod.test_creature.spin loop", "播放指定 Bedrock 动画", "通过"),
        ("停止并重置动画", "/magicmodel client animation stop target；/magicmodel client animation reset target", "停止并重置动画状态", "通过"),
        ("Tag 规则与 FBX 演示", "添加 magicmod_smoke_target 标签后调用 tag 替换规则", "验证标签目标命中并渲染 FBX 演示模型", "通过，FBX 演示计数 5604 -> 7010"),
        ("MMESH 中间格式替换", "/magicmodel client replace target mmesh magicmod:entity_models/runtime/test_cube.mmesh minecraft:textures/block/copper_block.png", "加载 .mmesh runtime 模型", "通过，MMESH 计数 7716 -> 9128"),
        ("MMESH fallback 动画", "/magicmodel client animation play target spin once", "播放 MMESH fallback 动画", "通过"),
        ("清理与重载释放", "/magicmodel client clear target + reload", "释放 runtime 和 GPU buffer", "通过，liveGpuBuffers=0"),
    ]
    for index, (title, command, expected, result) in enumerate(command_cases, start=1):
        add_command_case(document, index, title, command, expected, result)

    add_heading(document, "4. 游戏内日志证据", 1)
    add_code_block(document, [
        "[MODEL_SMOKE] Singleplayer world detected; starting model smoke sequence",
        "[MODEL_SMOKE] /magicmodel client replace target magicmod:test_cube",
        "[MODEL_SMOKE] Replacement rendered in-world; before=0, after=1386",
        "[MODEL_SMOKE] /magicmodel client replace target bedrock ...",
        "[MODEL_SMOKE] Bedrock replacement rendered in-world; before=2794, after=4198",
        "[MODEL_SMOKE] Tagged player and replaced tag with ASCII FBX demo model",
        "[MODEL_SMOKE] FBX replacement rendered in-world; before=5604, after=7010",
        "[MODEL_SMOKE] /magicmodel client replace target mmesh ...",
        "[MODEL_SMOKE] MMESH replacement rendered in-world; before=7716, after=9128",
        "[MODEL_SMOKE] Cache released cleanly: typeRules=1, runtimeInstances=0, loadedModels=0, liveGpuBuffers=0",
        "BUILD SUCCESSFUL in 5m 25s",
    ])

    add_heading(document, "5. 截图附图", 1)
    add_body(document, f"截图路径：{SCREENSHOT_PATH}")
    if SCREENSHOT_PATH.exists():
        document.add_picture(str(SCREENSHOT_PATH), width=Cm(15.5))
        caption = document.add_paragraph("图 1：MMESH runtime cube 替换玩家实体后的游戏内截图。")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(90, 90, 90)

    add_heading(document, "6. 结论与边界", 1)
    add_body(document, "本轮已跑通客户端本地模型替换主闭环：精确实体替换、Bedrock 模型和动画、FBX 演示导入、tag 规则、MMESH 中间格式、指令播放/停止/重置、reload 释放和退出清理。")
    add_body(document, "完整二进制 FBX 骨骼动画、服务端同步网络包、shader GPU skinning、动画混合树和 IK 仍应作为后续阶段实现。本轮没有把这些高风险能力伪装成已完成。")

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
